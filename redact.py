
"""

File: redact.py

Project: Automatic legal-document redactionn desktop application

Author: Ryan Friedman

Description: This program will allow a user to supply a Microsoft Word
             Document and a list of strings that need to be redacted.
             It will create a new Word Document with any version of the proper
             nouns replaced with a black highlight. We create a new document to
             ensure fresh metadata.

"""
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import tkinter as tk
from tkinter import *
from tkinter import filedialog
from string import punctuation


def binary_search(source, target):
    """ This function searches for a target in source in logarithmic time """

    left = 0
    right = len(source) - 1

    while left <= right:
        median = left + ((right - left) // 2)
        if source[median] == target:
            return True
        elif source[median] < target:
            left = median + 1
        else:
            right = median - 1

    return False


def popupmsg(msg):
    """ This function opens a pop-up window and displays a message. The user
        may click 'okay' to close the window. """

    popup = tk.Tk()
    popup.wm_title("Auto-Redaction App")
    label = tk.Label(popup, text=msg, font=("Helvetica", 10))
    label.pack(side="top", fill="x", pady=10)
    B1 = tk.Button(popup, text="Okay", command=popup.destroy)
    B1.pack()
    popup.mainloop()


def getDirFromFile(file_path):
    """ This function finds the last forward or back slash in directory path
        and slices the file name from it """

    break_index = -1
    for i in range(len(file_path) - 1, -1, -1):
        if file_path[i] == "\\" or file_path[i] == "/":
            break_index = i
            break

    if break_index != -1:
        file_path = file_path[:break_index + 1]

    return file_path


def processInfoFile(info_path):
    """ Opens the file and cleans/processes the information to be redacted
        into the proper state """

    redact_file = open(info_path, "r")
    redact_info = redact_file.read()
    redact_file.close()

    redact_info = redact_info.split(" ")

    for i in range(len(redact_info)):
        redact_info[i] = redact_info[i].strip(",")

    for entry in redact_info:
        if entry == "" or entry == "\n":
            redact_info.remove(entry)
        if " " in entry:
            entry = entry.strip()

    redact_info.sort()

    return redact_info


def processFiles():
    """ file_path_1 is our Word document and file_path_2 is any text file
        with proper nouns listed on separate lines. This is where we create
        our redacted Word File """

    # if word_file != "" and txt_file != "":
    #     popupmsg("You're file is being processed. You will receive another pop-up message when it is complete.")
    # else:
    #     exit(-1)

    file_path_1 = word_file
    file_path_2 = txt_file

    redact_info = processInfoFile(file_path_2)

    doc = Document(file_path_1)

    new_doc = Document()

    new_doc.settings.odd_and_even_pages_header_footer = doc.settings.odd_and_even_pages_header_footer

    temp_doc = Document()

    for i in range(len(doc.paragraphs)):
        processPara(doc.paragraphs[i], redact_info, new_doc, temp_doc)

    new_doc.save(getDirFromFile(file_path_1) + "redacted version.docx")
    popupmsg("Your redacted file has been created in the following directory: " + getDirFromFile(file_path_1))


def processPara(para, redact_info, new_doc, temp_doc):
    """ Gets indices for words that are to be redacts them then creates a
        a paragraph with those words redacted  """

    # gets indices within para.text of sensitive info
    redact_indices = getRedactIndices(para, redact_info)

    # creates a copy of para where each char is its own run
    para = convertRuns(para, temp_doc)

    redact(para, redact_indices, new_doc)


def getRedactIndices(para, redact_info):
    """ Returns a list containing tuples that represent the indices that any
        instance of redact_info spans in para """

    redact_indices = []

    curr_word = ""
    for i in range(len(para.text)):
        char = para.text[i]
        if char == " " or char == "\t" or char == "\n":

            if len(curr_word) >= 2:
                if curr_word[-2:] == "'s":
                    curr_word = char[:-2]

            prev_word = curr_word
            curr_word = curr_word.strip(punctuation)
            if binary_search(redact_info, curr_word):

                if curr_word != prev_word:
                    redact_indices.append((i - 1 - len(curr_word), i - 1))

                else:
                    redact_indices.append((i - len(curr_word), i))

            curr_word = ""

        else:
            curr_word += char

    # check the final word in each paragraph
    if curr_word:
        curr_word = curr_word.strip(punctuation)
        if binary_search(redact_info, curr_word):
            redact_indices.append((i - len(curr_word) + 1, len(para.text)))

    return redact_indices


def convertRuns(para, temp_doc):
    """ Converts all text in a paragraph to individual runs for each character
        so that we may uniformly redact instances of sensitive information."""

    p = temp_doc.add_paragraph()

    for run in para.runs:
        for char in run.text:
            r = p.add_run(char)
            r.bold = run.bold
            r.italic = run.italic
            r.underline = run.underline
            r.font.color.rgb = run.font.color.rgb
            r.font.name = run.font.name
            r.style.name = run.style.name
            r.font.size = run.font.size
            r.font.subscript = run.font.subscript
            r.font.superscript = run.font.superscript

    p.paragraph_format.alignment = para.paragraph_format.alignment
    p.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
    p.paragraph_format.keep_together = para.paragraph_format.keep_together
    p.paragraph_format.keep_with_next = para.paragraph_format.keep_with_next
    p.paragraph_format.left_indent = para.paragraph_format.left_indent
    p.paragraph_format.line_spacing = para.paragraph_format.line_spacing
    p.paragraph_format.line_spacing_rule = para.paragraph_format.line_spacing_rule
    p.paragraph_format.page_break_before = para.paragraph_format.page_break_before
    p.paragraph_format.right_indent = para.paragraph_format.right_indent
    p.paragraph_format.space_after = para.paragraph_format.space_after
    p.paragraph_format.widow_control = para.paragraph_format.widow_control
    p.paragraph_format.space_before = para.paragraph_format.space_before
    p.style = para.style

    return p


def redact(para, redact_indices, new_doc):
    """ This function modifies our paragraph object with new 'black runs'
        which are our redactions """

    index_map = getIndexMap(para, redact_indices)

    p = new_doc.add_paragraph()

    for i in range(len(para.runs)):
        if index_map[i] == 1:
            run = para.runs[i]
            r = p.add_run("X")
            r.font.highlight_color = WD_COLOR_INDEX.BLACK
            r.bold = run.bold
            r.italic = run.italic
            r.underline = run.underline
            r.font.color.rgb = run.font.color.rgb
            r.font.name = run.font.name
            r.style.name = run.style.name
            r.font.size = run.font.size
            r.font.subscript = run.font.subscript
            r.font.superscript = run.font.superscript
        else:
            run = para.runs[i]
            r = p.add_run(run.text)
            r.bold = run.bold
            r.italic = run.italic
            r.underline = run.underline
            r.font.color.rgb = run.font.color.rgb
            r.font.name = run.font.name
            r.style.name = run.style.name
            r.font.size = run.font.size
            r.font.subscript = run.font.subscript
            r.font.superscript = run.font.superscript

    p.paragraph_format.alignment = para.paragraph_format.alignment
    p.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
    p.paragraph_format.keep_together = para.paragraph_format.keep_together
    p.paragraph_format.keep_with_next = para.paragraph_format.keep_with_next
    p.paragraph_format.left_indent = para.paragraph_format.left_indent
    p.paragraph_format.line_spacing = para.paragraph_format.line_spacing
    p.paragraph_format.line_spacing_rule = para.paragraph_format.line_spacing_rule
    p.paragraph_format.page_break_before = para.paragraph_format.page_break_before
    p.paragraph_format.right_indent = para.paragraph_format.right_indent
    p.paragraph_format.space_after = para.paragraph_format.space_after
    p.paragraph_format.widow_control = para.paragraph_format.widow_control
    p.paragraph_format.space_before = para.paragraph_format.space_before
    p.style = para.style


def getIndexMap(para, redact_indices):
    """ creates a hash table that allows us to check in expected constant time
        whether a given character should be redacted or not
        (linear time to build) """

    index_map = {}

    for i in range(len(para.text)):
        index_map[i] = 0

    for entry in redact_indices:
        for i in range(entry[0], entry[1]):
            index_map[i] = 1

    return index_map


def requestFile(file_type):
    """ This function opens a file selection window using tkinter and
    assings the file path to the proper global variable """

    global word_file
    global txt_file

    root = Tk()
    root.withdraw()

    file_path = filedialog.askopenfilename()

    if not file_path:
        exit(1)

    if file_type == "docx":
        if file_path[-4:] != "docx":
            popupmsg("The file you selected was not a .docx file")

        else:
            word_file = file_path

    elif file_type == "txt":
        if file_path[-3:] != "txt":
            popupmsg("The file you selected was not a .txt file")
        else:
            txt_file = file_path


def request_docx():
    """ call to request file specifically for docx files"""
    requestFile("docx")


def request_txt():
    """ call to request file specifically for txt files"""
    requestFile("txt")


def on_closing():
    """ error code exit for closing the GUI"""

    exit(-1)


def GUI():
    """ Creates the GUI for this application. """

    global word_file
    global txt_file
    word_file = ""
    txt_file = ""

    master = tk.Tk()
    master.wm_title("Auto-redaction App")
    master.minsize(300, 200)
    master.geometry("300x200")

    b1 = Button(master,
                text="Click to add Word file",
                command=request_docx,
                height=0,
                width=0)
    b1.place(relx=0.5, rely=0.2, anchor=CENTER)

    b2 = Button(master,
                text="Click to add text file",
                command=request_txt,
                height=0,
                width=0)

    b2.place(relx=0.5, rely=0.5, anchor=CENTER)

    b3 = Button(master,
                text="Process",
                command=processFiles,
                height=0,
                width=0)

    b3.place(relx=0.5, rely=0.8, anchor=CENTER)

    master.protocol("WM_DELETE_WINDOW", on_closing)
    master.mainloop()


def main():

    GUI()


if __name__ == "__main__":
    main()
